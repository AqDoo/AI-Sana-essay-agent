"""
Генератор отчётов в формате Word (.docx)
и чтение эссе из Word файлов
"""

import io
from datetime import datetime
from typing import Dict, Optional

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def read_docx(file_bytes: bytes) -> str:
    """
    Читает текст из .docx файла.

    Args:
        file_bytes: байты загруженного файла

    Returns:
        Текст документа
    """
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())
    return "\n".join(paragraphs)


def generate_report(
    essay_text: str,
    quality: Optional[Dict],
    ai_result: Optional[Dict],
    plagiarism: Optional[Dict],
    student_name: str = ""
) -> bytes:
    """
    Генерирует отчёт в формате .docx.

    Args:
        essay_text: исходный текст эссе
        quality: результаты оценки качества
        ai_result: результаты AI детекции
        plagiarism: результаты проверки плагиата
        student_name: имя студента (опционально)

    Returns:
        bytes — содержимое .docx файла
    """
    doc = Document()

    # ── Настройка страницы ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # ── Стили ──
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # ── ЗАГОЛОВОК ──
    title = doc.add_heading('Отчёт оценки письменной работы', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # ── Мета-информация ──
    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_str = datetime.now().strftime('%d.%m.%Y %H:%M')
    if student_name:
        meta.add_run(f'Студент: {student_name}   |   ').font.size = Pt(10)
    run = meta.add_run(f'Дата проверки: {date_str}')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    _add_divider(doc)

    # ── БЛОК 1: Оценка качества ──
    if quality:
        doc.add_heading('Оценка качества текста', level=2)

        # Таблица оценок
        table = doc.add_table(rows=2, cols=5)
        table.style = 'Table Grid'

        headers = ['Грамматика', 'Структура', 'Содержание', 'Стиль', 'ИТОГО']
        scores = [
            f"{quality['grammar']}/25",
            f"{quality['structure']}/25",
            f"{quality['content']}/25",
            f"{quality['style']}/25",
            f"{quality['total']}/100"
        ]

        for i, (h, s) in enumerate(zip(headers, scores)):
            # Заголовок
            cell = table.rows[0].cells[i]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_bg(cell, 'D6E4F0')

            # Значение
            cell = table.rows[1].cells[i]
            cell.text = s
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(13)
            run.bold = True
            run.font.color.rgb = _score_color(
                quality['total'] if i == 4 else [
                    quality['grammar'], quality['structure'],
                    quality['content'], quality['style']
                ][i],
                100 if i == 4 else 25
            )
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Уровень
        doc.add_paragraph()
        grade_p = doc.add_paragraph()
        grade_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = grade_p.add_run(f"Уровень: {quality['grade']} ({quality['grade_letter']})")
        run.bold = True
        run.font.size = Pt(12)

        # Язык и статистика
        stats = doc.add_paragraph()
        stats.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lang = quality.get('language_detected', '?').upper()
        run = stats.add_run(
            f"Язык: {lang}  |  Слов: {quality['word_count']}  |  "
            f"Предложений: {quality['sentence_count']}  |  "
            f"Ср. длина предложения: {quality['avg_sentence_length']} слов"
        )
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

        doc.add_paragraph()

        # Комментарии по критериям
        doc.add_heading('Комментарии по критериям', level=3)
        criteria = [
            ('Грамматика', 'grammar_comment'),
            ('Структура', 'structure_comment'),
            ('Содержание', 'content_comment'),
            ('Стиль', 'style_comment'),
        ]
        for label, key in criteria:
            comment = quality.get(key, '')
            if comment:
                p = doc.add_paragraph()
                p.add_run(f'{label}: ').bold = True
                p.add_run(comment)

        doc.add_paragraph()

        # Сильные стороны и улучшения
        if quality.get('strengths') or quality.get('improvements'):
            cols_table = doc.add_table(rows=1, cols=2)
            cols_table.style = 'Table Grid'

            left_cell = cols_table.rows[0].cells[0]
            right_cell = cols_table.rows[0].cells[1]

            if quality.get('strengths'):
                p = left_cell.add_paragraph()
                p.add_run('Сильные стороны:').bold = True
                for s in quality['strengths']:
                    left_cell.add_paragraph(f'✓ {s}', style='List Bullet')

            if quality.get('improvements'):
                p = right_cell.add_paragraph()
                p.add_run('Что улучшить:').bold = True
                for imp in quality['improvements']:
                    right_cell.add_paragraph(f'→ {imp}', style='List Bullet')

        # Общий вывод
        if quality.get('general_feedback'):
            doc.add_paragraph()
            fb_para = doc.add_paragraph()
            fb_para.add_run('Общий вывод: ').bold = True
            fb_para.add_run(quality['general_feedback'])

    _add_divider(doc)

    # ── БЛОК 2: AI Детекция ──
    if ai_result:
        doc.add_heading('AI-детекция', level=2)

        ai_pct = ai_result.get('ai_percentage', 0)
        status = ai_result.get('status', 'unknown')

        # Статус
        status_p = doc.add_paragraph()
        verdict_run = status_p.add_run(f"{ai_result.get('verdict', '')}  ")
        verdict_run.bold = True
        verdict_run.font.size = Pt(12)
        if status == 'ai':
            verdict_run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        elif status == 'mixed':
            verdict_run.font.color.rgb = RGBColor(0xFF, 0x80, 0x00)
        else:
            verdict_run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)

        pct_run = status_p.add_run(f"({ai_pct:.1f}%)")
        pct_run.font.size = Pt(12)

        # Причины
        if ai_result.get('reasons'):
            doc.add_paragraph()
            doc.add_paragraph('Признаки анализа:').runs[0].bold = True
            for reason in ai_result['reasons']:
                doc.add_paragraph(f'• {reason}')

        # AI фразы
        if ai_result.get('ai_phrases_found'):
            p = doc.add_paragraph()
            p.add_run('Найденные AI-фразы: ').bold = True
            p.add_run(', '.join(ai_result['ai_phrases_found']))

    _add_divider(doc)

    # ── БЛОК 3: Плагиат ──
    if plagiarism:
        doc.add_heading('Проверка на плагиат', level=2)

        plag_pct = plagiarism.get('similarity_percentage', 0)
        plag_status = plagiarism.get('status', 'unknown')

        # Статус
        verdict_p = doc.add_paragraph()
        v_run = verdict_p.add_run(f"{plagiarism.get('verdict', '')}  ")
        v_run.bold = True
        v_run.font.size = Pt(12)
        if plag_status == 'plagiarized':
            v_run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        elif plag_status == 'suspicious':
            v_run.font.color.rgb = RGBColor(0xFF, 0x80, 0x00)
        else:
            v_run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)

        pct_run = verdict_p.add_run(f"({plag_pct:.1f}%)")
        pct_run.font.size = Pt(12)

        # Детали
        info_p = doc.add_paragraph()
        info_p.add_run(
            f"Проверено фраз: {plagiarism.get('phrases_checked', 0)}  |  "
            f"Совпадений: {plagiarism.get('phrases_matched', 0)}"
        ).font.color.rgb = RGBColor(0x70, 0x70, 0x70)

        # Найденные источники
        matches = plagiarism.get('matches', [])
        if matches:
            doc.add_paragraph()
            doc.add_paragraph('Найденные источники:').runs[0].bold = True

            for match in matches:
                p = doc.add_paragraph()
                p.add_run('Фраза: ').bold = True
                p.add_run(f'"{match.get("phrase", "")}"')

                p2 = doc.add_paragraph()
                p2.add_run('Источник: ').bold = True
                p2.add_run(match.get('source_title', match.get('source_url', '')))

                if match.get('source_url'):
                    p3 = doc.add_paragraph()
                    p3.add_run('URL: ').bold = True
                    url_run = p3.add_run(match['source_url'])
                    url_run.font.color.rgb = RGBColor(0x00, 0x56, 0xB3)

                doc.add_paragraph()

    _add_divider(doc)

    # ── ИСХОДНЫЙ ТЕКСТ ──
    doc.add_heading('Исходный текст эссе', level=2)
    for para_text in essay_text.split('\n'):
        if para_text.strip():
            p = doc.add_paragraph(para_text.strip())
            p.runs[0].font.size = Pt(10) if p.runs else None

    # ── Сохраняем в bytes ──
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def _score_color(score: int, max_score: int) -> RGBColor:
    """Возвращает цвет в зависимости от процента."""
    pct = score / max_score
    if pct >= 0.8:
        return RGBColor(0x00, 0x80, 0x00)  # зелёный
    elif pct >= 0.6:
        return RGBColor(0xFF, 0x80, 0x00)  # оранжевый
    else:
        return RGBColor(0xC0, 0x00, 0x00)  # красный


def _add_divider(doc: Document):
    """Добавляет горизонтальный разделитель."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_cell_bg(cell, hex_color: str):
    """Устанавливает цвет фона ячейки таблицы."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)
